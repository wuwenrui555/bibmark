"""
Write formatted citations to .docx, .md, and .tex files.
"""

from docx import Document


def write_docx(sections: list[tuple[str | None, list]], output_path: str, continuous_numbering: bool = False):
    """
    Write citations to a .docx file.

    Parameters
    ----------
    sections : list[tuple[str | None, list[list[Segment]]]]
        Each element is a ``(heading, citations)`` pair. ``heading`` is
        ``None`` for a flat (ungrouped) bibliography, or a section heading
        string for a grouped one. ``citations`` is a list of Segment lists.
    output_path : str
        Destination file path.
    continuous_numbering : bool, optional
        If ``True``, numbering continues across sections. If ``False``
        (default), numbering resets to 1 at the start of each section.
    """
    doc = Document()
    doc.add_heading("Bibliography", level=1)
    counter = 1
    for heading, citations in sections:
        if heading is not None:
            doc.add_heading(heading, level=2)
        if not continuous_numbering:
            counter = 1
        for segments in citations:
            para = doc.add_paragraph()
            para.add_run(f"{counter}. ")
            counter += 1
            for seg in segments:
                run = para.add_run(seg["text"])
                run.bold = seg["bold"]
                run.italic = seg["italic"]
                run.underline = seg["underline"]
                if seg["superscript"]:
                    run.font.superscript = True
    doc.save(output_path)


def write_md(sections: list[tuple[str | None, list[str]]], output_path: str, continuous_numbering: bool = False):
    """
    Write citations to a Markdown file.

    Parameters
    ----------
    sections : list[tuple[str | None, list[str]]]
        Each element is a ``(heading, citations)`` pair. ``heading`` is
        ``None`` for a flat (ungrouped) bibliography, or a section heading
        string for a grouped one. ``citations`` is a list of Markdown strings.
    output_path : str
        Destination file path.
    continuous_numbering : bool, optional
        If ``True``, numbering continues across sections. If ``False``
        (default), numbering resets to 1 at the start of each section.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("# Bibliography\n\n")
        counter = 1
        for heading, citations in sections:
            if heading is not None:
                f.write(f"## {heading}\n\n")
            if not continuous_numbering:
                counter = 1
            for citation in citations:
                f.write(f"{counter}. {citation}\n\n")
                counter += 1


def write_tex(sections: list[tuple[str | None, list[str]]], output_path: str, continuous_numbering: bool = False):
    """
    Write citations to a LaTeX file.

    Parameters
    ----------
    sections : list[tuple[str | None, list[str]]]
        Each element is a ``(heading, citations)`` pair. ``heading`` is
        ``None`` for a flat (ungrouped) bibliography, or a section heading
        string for a grouped one. ``citations`` is a list of LaTeX strings.
    output_path : str
        Destination file path.
    continuous_numbering : bool, optional
        If ``True``, numbering continues across sections. If ``False``
        (default), numbering resets to 1 at the start of each section.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\\documentclass{article}\n")
        f.write("\\usepackage[T1]{fontenc}\n")
        f.write("\\usepackage[utf8]{inputenc}\n")
        f.write("\\usepackage{hyperref}\n")
        f.write("\\setlength{\\parindent}{0pt}\n")
        f.write("\\begin{document}\n\n")
        f.write("\\section*{Bibliography}\n\n")
        counter = 1
        for heading, citations in sections:
            if heading is not None:
                f.write(f"\\subsection*{{{heading}}}\n\n")
            if not continuous_numbering:
                counter = 1
            for citation in citations:
                f.write(f"{counter}. {citation}\n\n")
                counter += 1
        f.write("\\end{document}\n")
